"""Generate Assignment3-Screenshots.docx with organized sections + placeholders for screenshots."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"H:\School\Labs\Springboot\Assignment3-Screenshots.docx"


def shade(cell, fill="F2F2F2"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def mono_run(p, text):
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    return run


def code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.25)
    mono_run(p, text)


def placeholder(doc, label):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = tbl.rows[0].cells[0]
    shade(cell, "FFF8DC")
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell.paragraphs[0].add_run(f"[ Paste screenshot here: {label} ]")
    run.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x66, 0x00)
    run.font.size = Pt(10)
    doc.add_paragraph()


def screenshot_item(doc, number, title, caption, *, method=None, path=None, body=None, expected=None, note=None):
    h = doc.add_heading(f"{number}  {title}", level=3)

    # caption
    p = doc.add_paragraph()
    run = p.add_run("Caption: ")
    run.bold = True
    p.add_run(caption)

    if method and path:
        p = doc.add_paragraph()
        run = p.add_run("Endpoint: ")
        run.bold = True
        mono_run(p, f"{method} {path}")

    if body:
        p = doc.add_paragraph()
        run = p.add_run("Request body:")
        run.bold = True
        code_block(doc, body)

    if expected:
        p = doc.add_paragraph()
        run = p.add_run("Expected: ")
        run.bold = True
        p.add_run(expected)

    if note:
        p = doc.add_paragraph()
        run = p.add_run("Note: ")
        run.bold = True
        run.italic = True
        note_run = p.add_run(note)
        note_run.italic = True

    placeholder(doc, f"{number} — {title}")


def section_header(doc, num, title, intro=None):
    doc.add_heading(f"Section {num} — {title}", level=1)
    if intro:
        doc.add_paragraph(intro)


def main():
    doc = Document()

    # Base styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---------- Cover ----------
    title = doc.add_heading("CPRG220 — Assignment 3: Security", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("EventHub API — Postman & Swagger Screenshots")
    r.bold = True
    r.font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Student: ").bold = True
    meta.add_run("_________________________    ")
    meta.add_run("Date: ").bold = True
    meta.add_run("_________________________")

    doc.add_paragraph()

    # ---------- Seeded accounts ----------
    doc.add_heading("Seeded Dev Accounts", level=2)
    doc.add_paragraph(
        "DevDataSeeder (dev profile) bootstraps two accounts on startup. "
        "Use these for all screenshots below."
    )
    t = doc.add_table(rows=3, cols=3)
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Role", "Email", "Password"
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    t.rows[1].cells[0].text = "ADMIN"
    t.rows[1].cells[1].text = "admin@eventhub.com"
    t.rows[1].cells[2].text = "admin123"
    t.rows[2].cells[0].text = "USER"
    t.rows[2].cells[1].text = "user@eventhub.com"
    t.rows[2].cells[2].text = "user123"

    doc.add_paragraph()

    # ---------- TOC ----------
    doc.add_heading("Contents", level=2)
    contents = [
        "1.  Authentication (register, login, invalid credentials)",
        "2.  Password Reset",
        "3.  Public Endpoints (no auth)",
        "4.  USER Role — Authorized",
        "5.  ADMIN Role — Authorized",
        "6.  RBAC Failure Cases (401 / 403)",
        "7.  Security Headers",
        "8.  Input Validation",
        "9.  Swagger Documentation",
        "10. Tests Passing",
        "11. Postman Runner Summary (bonus)",
    ]
    for line in contents:
        doc.add_paragraph(line, style="List Number" if False else None)

    doc.add_page_break()

    # ---------- Section 1 ----------
    section_header(doc, 1, "Authentication", "Four endpoints under /api/v1/auth. Tokens from login/register auto-populate Postman collection variables.")

    screenshot_item(doc, "1.1", "Register new user",
        "POST /api/v1/auth/register — creates a USER and returns a JWT.",
        method="POST", path="/api/v1/auth/register",
        body='{\n  "email": "newuser@example.com",\n  "firstName": "New",\n  "lastName": "User",\n  "password": "password123"\n}',
        expected="201 Created with { token, email, role: \"USER\" }")

    screenshot_item(doc, "1.2", "Login as USER",
        "POST /api/v1/auth/login — authenticates the seeded USER account.",
        method="POST", path="/api/v1/auth/login",
        body='{ "email": "user@eventhub.com", "password": "user123" }',
        expected="200 OK with JWT and role: \"USER\"")

    screenshot_item(doc, "1.3", "Login as ADMIN",
        "POST /api/v1/auth/login — authenticates the seeded ADMIN account.",
        method="POST", path="/api/v1/auth/login",
        body='{ "email": "admin@eventhub.com", "password": "admin123" }',
        expected="200 OK with JWT and role: \"ADMIN\"")

    screenshot_item(doc, "1.4", "Login with wrong password",
        "Invalid credentials rejected with 401 Unauthorized.",
        method="POST", path="/api/v1/auth/login",
        body='{ "email": "user@eventhub.com", "password": "wrongpassword" }',
        expected="401 Unauthorized")

    doc.add_page_break()

    # ---------- Section 2 ----------
    section_header(doc, 2, "Password Reset", "Two-step flow: request a token, then submit the token with a new password.")

    screenshot_item(doc, "2.1", "Request password reset",
        "POST /api/v1/auth/password-reset-request — generates a reset token (dev returns it in response body).",
        method="POST", path="/api/v1/auth/password-reset-request",
        body='{ "email": "user@eventhub.com" }',
        expected="200 OK with { resetToken, message }")

    screenshot_item(doc, "2.2", "Reset password with valid token",
        "POST /api/v1/auth/password-reset — consumes the token and sets a new password.",
        method="POST", path="/api/v1/auth/password-reset",
        body='{ "token": "<token from 2.1>", "newPassword": "newSecret123" }',
        expected="200 OK with success message")

    screenshot_item(doc, "2.3", "Reset password with invalid token",
        "Invalid or expired token rejected with 400 Bad Request.",
        method="POST", path="/api/v1/auth/password-reset",
        body='{ "token": "not-a-real-token", "newPassword": "whatever123" }',
        expected="400 Bad Request — Invalid or expired reset token")

    doc.add_page_break()

    # ---------- Section 3 ----------
    section_header(doc, 3, "Public Endpoints (no auth)", "Read-only event and category endpoints accessible without a JWT.")

    screenshot_item(doc, "3.1", "List events (public)",
        "GET /api/v1/events — paginated, no Authorization header.",
        method="GET", path="/api/v1/events",
        expected="200 OK with paginated events")

    screenshot_item(doc, "3.2", "Get event by ID (public)",
        "GET /api/v1/events/2 — anonymous read allowed.",
        method="GET", path="/api/v1/events/2",
        expected="200 OK with event detail")

    screenshot_item(doc, "3.3", "List categories (public)",
        "GET /api/v1/categories — no token required.",
        method="GET", path="/api/v1/categories",
        expected="200 OK with categories array")

    doc.add_page_break()

    # ---------- Section 4 ----------
    section_header(doc, 4, "USER Role — Authorized", "Requests sent with Authorization: Bearer {{userToken}}.")

    screenshot_item(doc, "4.1", "Create registration",
        "POST /api/v1/registrations — USER can book tickets.",
        method="POST", path="/api/v1/registrations",
        body='{\n  "items": [\n    { "eventId": 1, "quantity": 2 },\n    { "eventId": 3, "quantity": 1 }\n  ]\n}',
        expected="201 Created with registration + totalAmount")

    screenshot_item(doc, "4.2", "Get own registrations",
        "GET /api/v1/registrations/user/1 — returns the USER's own registrations.",
        method="GET", path="/api/v1/registrations/user/1",
        expected="200 OK with list of registrations")

    screenshot_item(doc, "4.3", "Cancel a registration",
        "PATCH /api/v1/registrations/1/cancel — status becomes CANCELLED.",
        method="PATCH", path="/api/v1/registrations/1/cancel",
        expected="200 OK with updated registration")

    doc.add_page_break()

    # ---------- Section 5 ----------
    section_header(doc, 5, "ADMIN Role — Authorized", "Requests sent with Authorization: Bearer {{adminToken}}.")

    screenshot_item(doc, "5.1", "Create category (ADMIN)",
        "POST /api/v1/categories — ADMIN-only write.",
        method="POST", path="/api/v1/categories",
        body='{ "name": "Workshops" }',
        expected="201 Created")

    screenshot_item(doc, "5.2", "Create event (ADMIN)",
        "POST /api/v1/events — ADMIN creates a new event.",
        method="POST", path="/api/v1/events",
        body='{\n  "name": "Spring Boot Workshop",\n  "description": "Learn Spring Boot fundamentals",\n  "ticketPrice": 25.00,\n  "categoryId": 2,\n  "eventDate": "2026-07-15T10:00:00"\n}',
        expected="201 Created")

    screenshot_item(doc, "5.3", "Update event (ADMIN)",
        "PUT /api/v1/events/2 — ADMIN updates existing event.",
        method="PUT", path="/api/v1/events/2",
        body='{\n  "name": "Spring Boot Workshop (Updated)",\n  "description": "Now with more annotations",\n  "ticketPrice": 35.00,\n  "categoryId": 2,\n  "eventDate": "2026-07-22T10:00:00"\n}',
        expected="200 OK")

    screenshot_item(doc, "5.4", "Delete event (ADMIN)",
        "DELETE /api/v1/events/5 — ADMIN removes an event.",
        method="DELETE", path="/api/v1/events/5",
        expected="204 No Content")

    screenshot_item(doc, "5.5", "List all registrations (ADMIN)",
        "GET /api/v1/registrations — ADMIN can view every registration.",
        method="GET", path="/api/v1/registrations",
        expected="200 OK with paginated list")

    doc.add_page_break()

    # ---------- Section 6 ----------
    section_header(doc, 6, "RBAC Failure Cases", "Proof that authentication and authorization rules are enforced.")

    screenshot_item(doc, "6.1", "No token on protected route",
        "POST /api/v1/events with NO Authorization header.",
        method="POST", path="/api/v1/events",
        expected="401 Unauthorized")

    screenshot_item(doc, "6.2", "Invalid token",
        "Authorization: Bearer not.a.real.token",
        method="GET", path="/api/v1/registrations",
        expected="401 Unauthorized")

    screenshot_item(doc, "6.3", "USER hitting ADMIN-only route",
        "POST /api/v1/events with USER token — authenticated but lacks role.",
        method="POST", path="/api/v1/events",
        expected="403 Forbidden")

    screenshot_item(doc, "6.4", "USER listing all registrations",
        "GET /api/v1/registrations with USER token — ADMIN-only endpoint.",
        method="GET", path="/api/v1/registrations",
        expected="403 Forbidden")

    screenshot_item(doc, "6.5", "USER deleting an event",
        "DELETE /api/v1/events/1 with USER token.",
        method="DELETE", path="/api/v1/events/1",
        expected="403 Forbidden")

    doc.add_page_break()

    # ---------- Section 7 ----------
    section_header(doc, 7, "Security Headers", "Evidence of HTTP header hardening configured in SecurityConfig.")

    screenshot_item(doc, "7.1", "Response headers on any endpoint",
        "Response Headers tab of any call (e.g. GET /api/v1/events) showing:\n"
        "  • X-Frame-Options: DENY\n"
        "  • X-Content-Type-Options: nosniff\n"
        "  • X-XSS-Protection: 1; mode=block\n"
        "  • Referrer-Policy: strict-origin-when-cross-origin\n"
        "  • Content-Security-Policy: default-src 'self'; frame-ancestors 'none'",
        note="In Postman, click the Headers tab on the response pane.")

    doc.add_page_break()

    # ---------- Section 8 ----------
    section_header(doc, 8, "Input Validation", "Bean Validation on auth DTOs (@NotBlank, @Email, @Size).")

    screenshot_item(doc, "8.1", "Register with short password",
        "Bean Validation rejects passwords shorter than 8 characters.",
        method="POST", path="/api/v1/auth/register",
        body='{ "email": "x@y.com", "firstName": "X", "lastName": "Y", "password": "short" }',
        expected="400 Bad Request with field error on password")

    screenshot_item(doc, "8.2", "Register with bad email",
        "@Email validation rejects malformed addresses.",
        method="POST", path="/api/v1/auth/register",
        body='{ "email": "not-an-email", "firstName": "X", "lastName": "Y", "password": "password123" }',
        expected="400 Bad Request with field error on email")

    doc.add_page_break()

    # ---------- Section 9 ----------
    section_header(doc, 9, "Swagger Documentation", "OpenAPI spec + interactive UI generated by SpringDoc.")

    screenshot_item(doc, "9.1", "Swagger UI — Authentication tag",
        "http://localhost:8080/swagger-ui.html — Authentication tag expanded, showing all four auth endpoints with request schemas.")

    screenshot_item(doc, "9.2", "Swagger UI — protected endpoint",
        "Swagger entry for POST /api/v1/events showing 401/403 responses and the request schema.")

    screenshot_item(doc, "9.3", "OpenAPI JSON",
        "http://localhost:8080/v3/api-docs — raw OpenAPI specification.")

    doc.add_page_break()

    # ---------- Section 10 ----------
    section_header(doc, 10, "Tests Passing", "Automated security tests confirm behavior end-to-end.")

    screenshot_item(doc, "10.1", "SecurityIntegrationTest output",
        "Terminal output from `./mvnw test -Dtest=SecurityIntegrationTest` — Tests run: 10, Failures: 0, Errors: 0.")

    screenshot_item(doc, "10.2", "JwtUtilTest output",
        "Terminal output from `./mvnw test -Dtest=JwtUtilTest` — all assertions pass.")

    doc.add_page_break()

    # ---------- Section 11 ----------
    section_header(doc, 11, "Postman Runner Summary (bonus)", "Full collection run exercising all five folders.")

    screenshot_item(doc, "11.1", "Collection Runner summary",
        "Postman → Collection → Run — summary panel showing pass counts across Auth, Public, USER role, ADMIN role, and RBAC failure cases.")

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
